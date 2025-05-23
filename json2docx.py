import json
import subprocess
import os
import sys 
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import qn as shared_qn
from datetime import datetime
import argparse

try:
    import jsonschema
    from jsonschema import validate, ValidationError
    JSONSCHEMA_AVAILABLE = True
except ImportError:
    JSONSCHEMA_AVAILABLE = False
    print("jsonschema not available - validation will be skipped")

class DocumentGenerator:
    def __init__(self):
        self.doc = None
        self.styles_defined = {}
        self.schema = self.get_json_schema()
    
    def get_json_schema(self):
        """Define JSON schema for document validation"""
        
        # Load from external JSON file:
        with open('schema/document_schema.json', 'r', encoding='utf-8') as f:
            document_schema = json.load(f)
        
        return document_schema

    def validate_json(self, json_data):
        """Validate JSON data against schema"""
        if not JSONSCHEMA_AVAILABLE:
            print("Warning: jsonschema not available, skipping validation")
            return True, "Validation skipped"
        
        try:
            validate(instance=json_data, schema=self.schema)
            return True, "Validation successful"
        except ValidationError as e:
            return False, str(e)
        except Exception as e:
            return False, f"Validation error: {str(e)}"

    def get_image_dimensions(self, image_path):
        """Get image dimensions using PIL"""
        try:
            with Image.open(image_path) as img:
                return {"width": img.width, "height": img.height}
        except Exception as e:
            print(f"Could not get dimensions for {image_path}: {e}")
            return {"width": 100, "height": 100}

    def add_field_code(self, paragraph, field_code, display_text):
        """Add a field code to a paragraph"""
        try:
            fld_begin = OxmlElement('w:fldSimple')
            fld_begin.set(qn('w:instr'), field_code)
            
            run = paragraph._element
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = display_text
            r.append(t)
            fld_begin.append(r)
            
            paragraph._element.append(fld_begin)
        except Exception as e:
            print(f"Field code creation failed: {e}")
            paragraph.add_run(display_text)

    def hex_to_rgb(self, hex_color):
        """Convert hex color to RGB tuple"""
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def apply_text_formatting(self, run, formatting):
        """Apply comprehensive text formatting"""
        if formatting.get('bold', False):
            run.bold = True
        if formatting.get('italic', False):
            run.italic = True
        if formatting.get('underline', False):
            run.underline = True
        if formatting.get('strike', False):
            run.font.strike = True
        if formatting.get('font_size'):
            run.font.size = Pt(formatting['font_size'])
        if formatting.get('font_name'):
            run.font.name = formatting['font_name']
        if formatting.get('color'):
            color = formatting['color']
            if isinstance(color, str) and color.startswith('#'):
                rgb = self.hex_to_rgb(color)
                run.font.color.rgb = RGBColor(*rgb)
        if formatting.get('highlight'):
            # Note: python-docx has limited highlight support
            pass
        if formatting.get('superscript'):
            run.font.superscript = True
        if formatting.get('subscript'):
            run.font.subscript = True

    def apply_paragraph_formatting(self, paragraph, formatting):
        """Apply paragraph-level formatting"""
        if formatting.get('alignment'):
            alignment_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(formatting['alignment'])
        
        if formatting.get('line_spacing'):
            spacing = formatting['line_spacing']
            if spacing == 'single':
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif spacing == 'double':
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif spacing == '1.5':
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif isinstance(spacing, (int, float)):
                paragraph.paragraph_format.line_spacing = spacing
        
        if formatting.get('space_before'):
            paragraph.paragraph_format.space_before = Pt(formatting['space_before'])
        if formatting.get('space_after'):
            paragraph.paragraph_format.space_after = Pt(formatting['space_after'])
        if formatting.get('left_indent'):
            paragraph.paragraph_format.left_indent = Inches(formatting['left_indent'])
        if formatting.get('right_indent'):
            paragraph.paragraph_format.right_indent = Inches(formatting['right_indent'])
        if formatting.get('first_line_indent'):
            paragraph.paragraph_format.first_line_indent = Inches(formatting['first_line_indent'])

    def create_custom_style(self, style_data):
        """Create custom paragraph or character styles"""
        style_name = style_data['name']
        style_type = style_data.get('type', 'paragraph')
        
        if style_name in self.styles_defined:
            return self.doc.styles[style_name]
        
        try:
            if style_type == 'paragraph':
                style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                if 'paragraph_format' in style_data:
                    # Apply paragraph formatting to style
                    para_format = style_data['paragraph_format']
                    if para_format.get('font_size'):
                        style.font.size = Pt(para_format['font_size'])
                    if para_format.get('font_name'):
                        style.font.name = para_format['font_name']
                    if para_format.get('bold'):
                        style.font.bold = para_format['bold']
                    if para_format.get('italic'):
                        style.font.italic = para_format['italic']
            
            elif style_type == 'character':
                style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                if 'font_format' in style_data:
                    font_format = style_data['font_format']
                    if font_format.get('font_size'):
                        style.font.size = Pt(font_format['font_size'])
                    if font_format.get('font_name'):
                        style.font.name = font_format['font_name']
                    if font_format.get('bold'):
                        style.font.bold = font_format['bold']
                    if font_format.get('italic'):
                        style.font.italic = font_format['italic']
            
            self.styles_defined[style_name] = True
            return style
        except Exception as e:
            print(f"Could not create style {style_name}: {e}")
            return None

    def create_table(self, table_data):
        """Create a table from JSON data"""
        rows = len(table_data['data'])
        cols = len(table_data['data'][0]) if rows > 0 else 0
        
        table = self.doc.add_table(rows=rows, cols=cols)
        
        # Apply table style if specified
        if table_data.get('style'):
            try:
                table.style = table_data['style']
            except:
                pass
        
        # Fill table data
        for row_idx, row_data in enumerate(table_data['data']):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                
                if isinstance(cell_data, dict):
                    # Complex cell with formatting
                    cell.text = cell_data.get('text', '')
                    if cell_data.get('formatting'):
                        # Apply cell formatting (limited support in python-docx)
                        pass
                else:
                    # Simple text cell
                    cell.text = str(cell_data)
        
        # Apply table formatting
        if table_data.get('formatting'):
            formatting = table_data['formatting']
            if formatting.get('autofit'):
                table.autofit = True

    def create_list(self, list_data):
        """Create numbered or bulleted lists"""
        list_type = list_data.get('list_type', 'bullet')  # 'bullet' or 'number'
        items = list_data.get('items', [])
        level = list_data.get('level', 0)
        
        for item in items:
            if isinstance(item, dict):
                # Complex list item
                text = item.get('text', '')
                formatting = item.get('formatting', {})
                
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                self.apply_text_formatting(run, formatting)
                
                # Set list style
                if list_type == 'number':
                    p.style = 'List Number'
                else:
                    p.style = 'List Bullet'
                    
                # Handle nested lists
                if item.get('subitems'):
                    nested_list = {
                        'type': 'list',
                        'list_type': list_type,
                        'items': item['subitems'],
                        'level': level + 1
                    }
                    self.create_list(nested_list)
            else:
                # Simple text item
                p = self.doc.add_paragraph(str(item))
                if list_type == 'number':
                    p.style = 'List Number'
                else:
                    p.style = 'List Bullet'

    def add_page_break(self):
        """Add a page break"""
        self.doc.add_page_break()

    def create_header_footer(self, section_data):
        """Create headers and footers"""
        if not section_data:
            return
            
        section = self.doc.sections[0]  # Assume first section for now
        
        # Header
        if section_data.get('header'):
            header = section.header
            header_p = header.paragraphs[0]
            header_p.text = section_data['header']['text']
            if section_data['header'].get('alignment'):
                alignment_map = {
                    'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                    'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                    'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
                }
                header_p.alignment = alignment_map.get(section_data['header']['alignment'])
        
        # Footer
        if section_data.get('footer'):
            footer = section.footer
            footer_p = footer.paragraphs[0]
            footer_p.text = section_data['footer']['text']
            if section_data['footer'].get('alignment'):
                alignment_map = {
                    'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                    'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                    'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
                }
                footer_p.alignment = alignment_map.get(section_data['footer']['alignment'])

    def set_page_layout(self, layout_data):
        """Set page layout properties"""
        if not layout_data:
            return
            
        section = self.doc.sections[0]
        
        if layout_data.get('orientation'):
            if layout_data['orientation'] == 'landscape':
                section.orientation = WD_ORIENTATION.LANDSCAPE
                # Swap width and height for landscape
                section.page_width, section.page_height = section.page_height, section.page_width
        
        if layout_data.get('margins'):
            margins = layout_data['margins']
            if margins.get('top'):
                section.top_margin = Inches(margins['top'])
            if margins.get('bottom'):
                section.bottom_margin = Inches(margins['bottom'])
            if margins.get('left'):
                section.left_margin = Inches(margins['left'])
            if margins.get('right'):
                section.right_margin = Inches(margins['right'])

    def create_paragraph_from_json(self, para_data):
        """Create a paragraph from JSON data with enhanced capabilities"""
        para_type = para_data.get('type', 'normal')
        
        # Handle special paragraph types
        if para_type == 'page_break':
            self.add_page_break()
            return
        elif para_type == 'table':
            self.create_table(para_data)
            return
        elif para_type == 'list':
            self.create_list(para_data)
            return
        
        if para_type == 'title_paragraph': # New explicit type for 'Title' style
            # doc.add_heading() with level=0 applies the built-in 'Title' style.
            # An empty paragraph with this style is created.
            # Text content will be added by the loop over 'content' items below.
            paragraph = self.doc.add_heading(level=0)
        elif para_type.startswith('heading'): # Handles "heading", "heading0", "heading1", etc.
            level_str = para_type.replace('heading', '').strip()
            level = 1 # Default level for a generic "heading" type or if parsing fails

            if level_str.isdigit(): # Handles "heading0", "heading1", "heading2", etc.
                level = int(level_str)
            elif 'level' in para_data: # Allows JSON like: { "type": "heading", "level": 0 }
                try:
                    level = int(para_data['level'])
                except (ValueError, TypeError):
                    print(f"Warning: Invalid 'level' value ('{para_data['level']}') for heading. Using default level 1.")
            # If para_type is just "heading" and no 'level' key, it defaults to level 1.
            paragraph = self.doc.add_heading(level=level)

        elif para_type == 'custom_style' and para_data.get('style_name'):
            paragraph = self.doc.add_paragraph()
            try:
                paragraph.style = para_data['style_name']
            except:
                print(f"Style {para_data['style_name']} not found")
        else:
            paragraph = self.doc.add_paragraph()
        
        # Apply paragraph formatting
        if para_data.get('formatting'):
            self.apply_paragraph_formatting(paragraph, para_data['formatting'])
        
        # Process content
        original_content = para_data.get('content', [])
        new_content = []
        i = 0
        n = len(original_content)

        while i < n:
            item_i = original_content[i]
            merged = False

            # Check for Three-Part Merge
            if i + 2 < n:
                item_i_plus_1 = original_content[i+1]
                item_i_plus_2 = original_content[i+2]

                # Ensure all are text types and text key exists
                if (item_i.get('type') == 'text' and item_i_plus_1.get('type') == 'text' and item_i_plus_2.get('type') == 'text' and
                    isinstance(item_i.get('text'), str) and isinstance(item_i_plus_1.get('text'), str) and isinstance(item_i_plus_2.get('text'), str)):
                    
                    text_i = item_i['text']
                    text_i_plus_1 = item_i_plus_1['text']
                    text_i_plus_2 = item_i_plus_2['text']
                    format_i = item_i.get('formatting', {})
                    format_i_plus_1 = item_i_plus_1.get('formatting', {})
                    format_i_plus_2 = item_i_plus_2.get('formatting', {})

                    if text_i.startswith('{') and text_i_plus_2 == '}' and format_i == format_i_plus_2:
                        merged_text = text_i[1:] + text_i_plus_1
                        merged_formatting = {**format_i, **format_i_plus_1} # Inner (i+1) takes precedence for shared keys
                        
                        new_content.append({
                            "type": "text",
                            "text": merged_text,
                            "formatting": merged_formatting
                        })
                        i += 3
                        merged = True
            
            # Check for Two-Part Merge (if three-part didn't match)
            if not merged and i + 1 < n:
                item_i_plus_1 = original_content[i+1]

                if (item_i.get('type') == 'text' and item_i_plus_1.get('type') == 'text' and
                    isinstance(item_i.get('text'), str) and isinstance(item_i_plus_1.get('text'), str)):

                    text_i = item_i['text']
                    text_i_plus_1 = item_i_plus_1['text']
                    format_i = item_i.get('formatting', {})
                    format_i_plus_1 = item_i_plus_1.get('formatting', {})

                    if text_i.startswith('{') and text_i_plus_1 == '}' and format_i == format_i_plus_1:
                        merged_text = text_i[1:]
                        merged_formatting = format_i
                        
                        if merged_text: # Only append if there's actual text after removing '{'
                            new_content.append({
                                "type": "text",
                                "text": merged_text,
                                "formatting": merged_formatting
                            })
                        i += 2
                        merged = True

            # Default Action
            if not merged:
                new_content.append(item_i)
                i += 1
        
        # Replace original content with preprocessed content
        content = new_content
        
        for item in content:
            try:
                if item['type'] == 'text':
                    run = paragraph.add_run(item['text'])
                    if 'formatting' in item:
                        self.apply_text_formatting(run, item['formatting'])
                        
                elif item['type'] == 'citation':
                    field_code = f"ADDIN ZOTERO_ITEM CSL_CITATION {json.dumps(item['field_data'])}"
                    display_text = item.get('display_text', '')
                    self.add_field_code(paragraph, field_code, display_text)
                    
                elif item['type'] == 'bibliography':
                    field_data = item.get('field_data', {"uncited": [], "omitted": [], "custom": []})
                    field_code = f"ADDIN ZOTERO_BIBL {json.dumps(field_data)} CSL_BIBLIOGRAPHY"
                    display_text = item.get('display_text', 'Bibliography')
                    self.add_field_code(paragraph, field_code, display_text)
                    
                elif item['type'] == 'image':
                    self.add_image_to_paragraph(paragraph, item)
                    
                elif item['type'] == 'field':
                    # Generic field support
                    field_code = item.get('field_code', '')
                    display_text = item.get('display_text', '')
                    self.add_field_code(paragraph, field_code, display_text)
                    
            except Exception as e:
                print(f"Error processing content item: {e}")
                paragraph.add_run(f"[Error: {str(e)}]")

    def add_image_to_paragraph(self, paragraph, image_data):
        """Add image with enhanced options"""
        try:
            image_path = image_data['path']
            
            # Validate image exists
            if not os.path.exists(image_path):
                paragraph.add_run(f"[Image not found: {image_path}]")
                return
            
            width = image_data.get('width_inches', 2.0)
            height = image_data.get('height_inches', 2.0)
            
            # Handle aspect ratio preservation
            if image_data.get('preserve_aspect_ratio', False):
                dimensions = self.get_image_dimensions(image_path)
                aspect_ratio = dimensions['width'] / dimensions['height']
                if width and not height:
                    height = width / aspect_ratio
                elif height and not width:
                    width = height * aspect_ratio
            
            run = paragraph.add_run()
            picture = run.add_picture(image_path, width=Inches(width), height=Inches(height))
            
            # Add caption if specified
            if image_data.get('caption'):
                caption_para = self.doc.add_paragraph()
                caption_run = caption_para.add_run(image_data['caption'])
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if image_data.get('caption_formatting'):
                    self.apply_text_formatting(caption_run, image_data['caption_formatting'])
                    
        except Exception as e:
            print(f"Could not add image {image_data.get('path', 'unknown')}: {e}")
            paragraph.add_run(f"[Image: {image_data.get('alt_text', 'Image could not be loaded')}]")

    def create_document_from_json(self, json_data):
        """Create a Word document from JSON configuration with validation"""
        try:
            # Validate JSON against schema
            is_valid, error_message = self.validate_json(json_data)
            if not is_valid:
                print(f"JSON validation failed: {error_message}")
                print("Exiting due to validation errors.")
                sys.exit(1) # Exit the script with a non-zero status code (error)
            else:
                print("JSON validation passed ✓")
            
            # Initialize document
            template_path = json_data.get('template_path')
            if template_path and os.path.exists(template_path):
                self.doc = Document(template_path)
            else:
                self.doc = Document()
            
            # Set document properties
            properties = json_data.get('properties', {})
            if 'title' in properties:
                self.doc.core_properties.title = properties['title']
            if 'author' in properties:
                self.doc.core_properties.author = properties['author']
            if 'subject' in properties:
                self.doc.core_properties.subject = properties['subject']
            if 'created' in properties:
                self.doc.core_properties.created = datetime.now()
            
            # Create custom styles
            for style_data in json_data.get('custom_styles', []):
                self.create_custom_style(style_data)
            
            # Set page layout
            self.set_page_layout(json_data.get('page_layout'))
            
            # Create header/footer
            self.create_header_footer(json_data.get('header_footer'))
            
            # Process content
            for para_data in json_data.get('content', []):
                self.create_paragraph_from_json(para_data)
            
            return self.doc
            
        except Exception as e:
            print(f"Error creating document: {e}")
            return None

def main():
    """Main function to generate document"""
    # --- START OF MODIFICATIONS ---
    parser = argparse.ArgumentParser(description="Generate a DOCX document from a JSON configuration.")
    parser.add_argument(
        "output_file", 
        help="Path to the output DOCX file. Folders will be created if they don't exist."
    )
    parser.add_argument(
        "--config",
        default="data/document_config.json",
        help="Path to the JSON configuration file (default: document_config.json)"
    )
    args = parser.parse_args()

    output_filename = args.output_file
    config_filepath = args.config

    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_filename)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created directory: {output_dir}")
    # --- END OF MODIFICATIONS ---

    # Initialize generator
    generator = DocumentGenerator()
    
    # Load from the specified JSON config file
    try:
        with open(config_filepath, 'r', encoding='utf-8') as f:
            document_config = json.load(f)
    except FileNotFoundError:
        print(f"Error: Configuration file '{config_filepath}' not found.")
        return
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from '{config_filepath}'.")
        return
    
    # Create document from JSON
    doc = generator.create_document_from_json(document_config)
    
    if doc:
        # Save the document
        doc.save(output_filename) # Use the CLI provided filename
        print(f"Document saved as {output_filename}")
        
    else:
        print("Failed to create document")


if __name__ == "__main__":
    main()