import argparse
import json
import os
import re
import sys # Added for sys.exit()
import uuid

import docx
from docx.shared import Cm, Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Compiled regexes for citation commands - Using raw strings
CITET_RE = re.compile(r'\\citet(?:\[(.*?)\])?\{(.*?)\}')
CITEP_RE = re.compile(r'\\citep(?:\[(.*?)\])?\{(.*?)\}')
CITEAUTHOR_RE = re.compile(r'\\citeauthor\{(.*?)\}')
CITEYEARPAR_RE = re.compile(r'\\citeyearpar(?:\[(.*?)\])?\{(.*?)\}')

CITATION_COMMANDS = [
    ('citet', CITET_RE),
    ('citep', CITEP_RE),
    ('citeauthor', CITEAUTHOR_RE),
    ('citeyearpar', CITEYEARPAR_RE),
]

def _apply_basic_formatting_to_str(text_segment: str) -> str:
    # This function is being deprecated by the new process_inline_elements logic
    # but kept for now to avoid breaking other parts if they call it.
    # New logic directly handles bold/italic within process_inline_elements.
    if not text_segment: return [] 
    parts = re.split(r'(\\textbf\{.*?\})|(\\textit\{.*?\})', text_segment)
    processed_parts = []
    for part in parts:
        if not part: continue
        if part.startswith('\\textbf{') and part.endswith('}'):
            processed_parts.append({'type': 'text', 'value': '**' + part[8:-1] + '**'})
        elif part.startswith('\\textit{') and part.endswith('}'):
            processed_parts.append({'type': 'text', 'value': '*' + part[8:-1] + '*'})
        else:
            processed_parts.append({'type': 'text', 'value': part})
    return processed_parts

def process_inline_elements(text_line: str) -> list:
    # New logic: directly parse text_line for text, bold, italic, and citations
    final_elements = []
    current_pos = 0

    # Regex to find \textbf{}, \textit{}, or any citation command
    # Non-greedy match for content within braces: (.*?)
    # Citation commands are already defined in CITATION_COMMANDS
    
    # Construct a combined regex pattern for all known inline commands
    # Order matters if commands can be substrings of each other (not typical for these)
    inline_patterns = []
    # Add textbf and textit first
    inline_patterns.append(r'(\\textbf\{(.*?)\})')
    inline_patterns.append(r'(\\textit\{(.*?)\})')
    # Add citation commands
    for _, cmd_re in CITATION_COMMANDS:
        inline_patterns.append(f'({cmd_re.pattern})') # Use .pattern to get the string regex

    # Combine all patterns with OR operator
    combined_pattern_str = "|".join(inline_patterns)
    tokenizer_re = re.compile(combined_pattern_str)

    while current_pos < len(text_line):
        match = tokenizer_re.search(text_line, current_pos)

        if not match:
            # No more special commands, the rest is plain text
            if current_pos < len(text_line):
                final_elements.append({'type': 'text', 'value': text_line[current_pos:]})
            break

        # Add preceding text if any
        if match.start() > current_pos:
            final_elements.append({'type': 'text', 'value': text_line[current_pos:match.start()]})
        
        # Determine which group matched to identify the command
        # The combined regex will have many groups. We need to find which one captured.
        # Group 0 is the whole match. Group 1 is \textbf content, Group 2 is its content.
        # Group 3 is \textit content, Group 4 is its content.
        # Then citation groups follow.
        
        # Check \textbf (groups 1 and 2)
        if match.group(1): # Full \textbf{content} match
            final_elements.append({'type': 'text', 'value': f"**{match.group(2)}**"}) # Content is group 2
        # Check \textit (groups 3 and 4)
        elif match.group(3): # Full \textit{content} match
            final_elements.append({'type': 'text', 'value': f"*{match.group(4)}*"}) # Content is group 4
        else:
            # Must be a citation. Iterate through CITATION_COMMANDS to find which one.
            # The match object `match` is from `tokenizer_re`. We need to re-match with individual citation regexes
            # on the matched segment to correctly parse groups for prenotes/postnotes.
            # This is because the group indices from tokenizer_re are hard to map back.
            
            citation_matched_here = False
            # Check which citation command regex matches at `match.start()`
            for cmd_name, cmd_re_obj in CITATION_COMMANDS:
                # Use match.group(0) which is the full matched segment by tokenizer_re
                # then try to match this segment with specific citation regex
                citation_specific_match = cmd_re_obj.fullmatch(match.group(0)) # Use fullmatch on the segment
                if citation_specific_match:
                    prenote = ""
                    keys_str = ""
                    if cmd_name == 'citeauthor':
                        keys_str = citation_specific_match.group(1)
                    else: 
                        prenote = citation_specific_match.group(1) if citation_specific_match.group(1) else ""
                        keys_str = citation_specific_match.group(2)
                    
                    parsed_keys = [k.strip() for k in keys_str.split(',')]
                    final_elements.append({
                        'type': 'citation',
                        'command': cmd_name,
                        'keys': parsed_keys,
                        'prenote': prenote.strip(),
                        'postnote': "" 
                    })
                    citation_matched_here = True
                    break # Found the specific citation type
            if not citation_matched_here:
                # Should not happen if tokenizer_re is correct and includes all citation patterns
                # But as a fallback, treat unrecognized match as text
                final_elements.append({'type': 'text', 'value': match.group(0)})

        current_pos = match.end()
        
    return [elem for elem in final_elements if elem.get('value') or elem.get('type') == 'citation']

def parse_latex_options(options_str: str) -> dict:
    options = {}
    if not options_str: return options
    for part in options_str.split(','):
        match = re.match(r'\s*(.*?)\s*=\s*(.*)\s*', part)
        if match:
            key, value = match.groups()
            options[key.strip()] = value.strip()
    return options

def parse_latex_to_json(latex_content: str) -> list:
    json_output = []
    lines = latex_content.splitlines()
    current_paragraph_text_accumulator = [] 
    in_itemize_env = False
    in_enumerate_env = False
    # Accumulator for list items, as they might span multiple lines if not handled by \item
    current_list_item_accumulator = [] 
    
    def flush_paragraph():
        nonlocal current_paragraph_text_accumulator, json_output
        if current_paragraph_text_accumulator:
            full_paragraph_str = " ".join(current_paragraph_text_accumulator).strip()
            if full_paragraph_str:
                json_output.append({'type': 'paragraph', 'content': process_inline_elements(full_paragraph_str)})
            current_paragraph_text_accumulator = []

    def flush_list_item():
        nonlocal current_list_item_accumulator, json_output, in_itemize_env, in_enumerate_env
        if current_list_item_accumulator:
            full_item_str = " ".join(current_list_item_accumulator).strip()
            if full_item_str:
                list_type = 'list_item_bullet' if in_itemize_env else 'list_item_ordered'
                json_output.append({'type': list_type, 'content': process_inline_elements(full_item_str)})
            current_list_item_accumulator = []

    env_item_pattern = re.compile(
        r'(\\begin\{(itemize|enumerate)\})|'
        r'(\\end\{(itemize|enumerate)\})|'
        r'(\\item(?:\s+|$))'  # \item followed by space or end of string
    )
    # Standalone command patterns (ensure they are checked on a line-by-line basis)
    title_pattern = re.compile(r'\\title\{(.*)\}')
    section_pattern = re.compile(r'\\section\{(.*)\}')
    subsection_pattern = re.compile(r'\\subsection\{(.*)\}')
    includegraphics_pattern = re.compile(r'\\includegraphics(?:\[(.*?)\])?\{(.*?)\}')

    # Process content line by line primarily for standalone commands,
    # then segment by segment for environments within those lines.
    input_lines = latex_content.splitlines()
    line_idx = 0

    while line_idx < len(input_lines):
        current_line_stripped = input_lines[line_idx].strip()
        line_idx += 1 # Consume line

        if not current_line_stripped: # Handle empty or whitespace-only lines
            # An empty line flushes current item (if multi-line) and current paragraph.
            # It does not necessarily end the list environment itself.
            flush_list_item()
            flush_paragraph()
            continue

        # Check for standalone commands that occupy the entire line
        title_match = title_pattern.fullmatch(current_line_stripped)
        section_match = section_pattern.fullmatch(current_line_stripped)
        subsection_match = subsection_pattern.fullmatch(current_line_stripped)
        image_match = includegraphics_pattern.fullmatch(current_line_stripped)

        if title_match:
            flush_paragraph(); flush_list_item() # Ensure previous context is cleared
            title_str = title_match.group(1).strip()
            content_elements = process_inline_elements(title_str)
            concatenated_content = "".join(el.get('value', '') for el in content_elements)
            json_output.append({'type': 'title', 'content': concatenated_content})
            continue
        elif section_match:
            flush_paragraph(); flush_list_item()
            section_str = section_match.group(1).strip()
            content_elements = process_inline_elements(section_str)
            concatenated_content = "".join(el.get('value', '') for el in content_elements)
            json_output.append({'type': 'heading1', 'content': concatenated_content})
            continue
        elif subsection_match:
            flush_paragraph(); flush_list_item()
            subsection_str = subsection_match.group(1).strip()
            content_elements = process_inline_elements(subsection_str)
            concatenated_content = "".join(el.get('value', '') for el in content_elements)
            json_output.append({'type': 'heading2', 'content': concatenated_content})
            continue
        elif image_match:
            flush_paragraph(); flush_list_item()
            options_str = image_match.group(1)
            image_path = image_match.group(2).strip()
            options_dict = parse_latex_options(options_str if options_str else "")
            json_output.append({'type': 'image', 'path': image_path, 'options': options_dict})
            continue

        # If not a standalone command fully matching the line,
        # process the line for environments, items, and text segments.
        processed_upto_in_line = 0
        while processed_upto_in_line < len(current_line_stripped):
            segment_to_parse = current_line_stripped[processed_upto_in_line:]
            match = env_item_pattern.search(segment_to_parse)
            
            text_before_command = segment_to_parse[:match.start()] if match else segment_to_parse
            
            if text_before_command:
                if in_itemize_env or in_enumerate_env:
                    current_list_item_accumulator.append(text_before_command)
                else:
                    current_paragraph_text_accumulator.append(text_before_command)
                processed_upto_in_line += len(text_before_command)
            
            if not match: break # No more env/item commands in this segment

            command_full_match = match.group(0)
            begin_env_match = match.group(1); env_type_begin = match.group(2)
            end_env_match = match.group(3); env_type_end = match.group(4) # For \end{...}
            item_cmd_match = match.group(5)

            if begin_env_match:
                flush_paragraph(); flush_list_item() # Clear context before starting list
                if env_type_begin == 'itemize': in_itemize_env = True
                elif env_type_begin == 'enumerate': in_enumerate_env = True
            elif end_env_match:
                flush_list_item() # Finalize last item
                if env_type_end == 'itemize' and in_itemize_env: in_itemize_env = False
                elif env_type_end == 'enumerate' and in_enumerate_env: in_enumerate_env = False
                # Do not flush_paragraph here, as text might follow on the same line after \end{env}
            elif item_cmd_match:
                if in_itemize_env or in_enumerate_env: flush_list_item() # Finalize previous item
                else: current_paragraph_text_accumulator.append(command_full_match) # \item outside list
            
            processed_upto_in_line += len(command_full_match)
        
    # After all lines are processed
    flush_list_item() 
    flush_paragraph() 
    return json_output

# --- DOCX Generation --- (Assumed correct from previous steps)
def construct_zotero_json_payload(citation_obj: dict) -> tuple[str, str]:
    citation_id = str(uuid.uuid4())
    keys = citation_obj.get('keys', [])
    prenote = citation_obj.get('prenote', '')
    command = citation_obj.get('command', 'citep') 

    plain_citation_parts = []
    keys_display_str = ", ".join(keys) 
    
    if command == 'citet':
        display_text = keys_display_str 
        if prenote:
            display_text += f" ({prenote})" 
    elif command == 'citeauthor':
        display_text = keys_display_str
    else: 
        if prenote:
            plain_citation_parts.append(prenote)
        plain_citation_parts.append(keys_display_str)
        display_text = f"({'; '.join(plain_citation_parts)})"
    
    plain_citation_for_display = display_text

    citation_items = []
    for key_idx, key in enumerate(keys):
        item_data = {"id": key, "type": "book"} 
        current_item_prenote = ""
        if key_idx == 0 and prenote:
             current_item_prenote = prenote
        
        citation_items.append({
            "itemData": item_data, 
            "prefix": current_item_prenote, 
            "suffix": "", 
            "uris": [f"http://zotero.org/users/local/placeholder/items/{key}"]
        })

    zotero_payload = {
        "citationID": citation_id,
        "properties": {"plainCitation": plain_citation_for_display },
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json"
    }
    return json.dumps(zotero_payload), plain_citation_for_display


def add_runs_for_formatted_text(paragraph_or_heading, text_with_markers: str):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text_with_markers) # Using raw string
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            paragraph_or_heading.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*'):
            paragraph_or_heading.add_run(part[1:-1]).italic = True
        else:
            paragraph_or_heading.add_run(part)

def generate_docx(json_data: list, output_docx_path: str):
    doc = docx.Document()
    doc.core_properties.title = "Converted LaTeX Document" 

    for item in json_data:
        item_type = item.get('type')
        
        if item_type == 'title':
            title_text = item.get('content', 'Untitled Document')
            doc.core_properties.title = title_text 
            h = doc.add_heading(level=0) 
            add_runs_for_formatted_text(h, title_text)
            
        elif item_type == 'heading1':
            h = doc.add_heading(level=1)
            add_runs_for_formatted_text(h, item.get('content', ''))
            
        elif item_type == 'heading2':
            h = doc.add_heading(level=2)
            add_runs_for_formatted_text(h, item.get('content', ''))
            
        elif item_type == 'paragraph' or item_type == 'list_item_bullet' or item_type == 'list_item_ordered':
            style = None
            if item_type == 'list_item_bullet': style = 'ListBullet'
            elif item_type == 'list_item_ordered': style = 'ListNumber'
            p = doc.add_paragraph(style=style)

            for content_element in item.get('content', []):
                element_type = content_element.get('type')
                if element_type == 'text':
                    add_runs_for_formatted_text(p, content_element.get('value', ''))
                elif element_type == 'citation':
                    zotero_json_string, plain_citation_text_for_display = construct_zotero_json_payload(content_element)
                    instr_text = f" ADDIN ZOTERO_ITEM CSL_CITATION {zotero_json_string}"
                    run = p.add_run()
                    fldSimple = OxmlElement('w:fldSimple')
                    fldSimple.set(qn('w:instr'), instr_text)
                    sub_r = OxmlElement('w:r')
                    sub_t = OxmlElement('w:t')
                    sub_t.text = plain_citation_text_for_display 
                    sub_r.append(sub_t)
                    fldSimple.append(sub_r)
                    run._r.append(fldSimple)

        elif item_type == 'image':
            image_path = item.get('path')
            options = item.get('options', {})
            if not os.path.exists(image_path):
                print(f"Warning: Image not found at {image_path}. Skipping.", file=sys.stderr)
                p = doc.add_paragraph()
                add_runs_for_formatted_text(p, f"[Image not found: {image_path}]")
                continue
            try:
                width_str = options.get('width')
                width_val = None
                if width_str:
                    val_match = re.match(r'([\d\.]+)\s*(cm|in|px)?', width_str) # Using raw string
                    if val_match: # This means value and potentially unit were matched
                        val = float(val_match.group(1))
                        unit = val_match.group(2) # This can be None if no unit like 'cm' or 'in' was present
                        if unit == 'cm': width_val = Cm(val)
                        elif unit == 'in': width_val = Inches(val)
                        elif unit is None and r'\textwidth' in width_str: # e.g. width=0.8\textwidth
                            print(f"Warning: Relative width '{width_str}' for image {image_path} not supported. Using default fallback width of 6 inches.", file=sys.stderr)
                            width_val = Inches(6)
                        elif unit is None: # Matched a number but no recognized unit and not textwidth
                            print(f"Warning: Width '{width_str}' for image {image_path} has an unrecognized or missing unit. Adding image with original dimensions.", file=sys.stderr)
                            # width_val remains None, image added with original dimensions
                        else: # unit is not None and not 'cm' or 'in'
                            print(f"Warning: Width unit '{unit}' for image {image_path} not directly supported. Adding image with original dimensions.", file=sys.stderr)
                            # width_val remains None
                    elif r'\textwidth' in width_str: # No specific value like '0.8' was matched, but textwidth is present (e.g. width=\textwidth)
                        print(f"Warning: Relative width '{width_str}' for image {image_path} not supported. Using default fallback width of 6 inches.", file=sys.stderr)
                        width_val = Inches(6)
                    else: 
                        print(f"Warning: Could not parse width '{width_str}' for image {image_path}. Adding image with original dimensions.", file=sys.stderr)
                        # width_val remains None
                
                if width_val: doc.add_picture(image_path, width=width_val)
                else: doc.add_picture(image_path)
            except Exception as e:
                print(f"Error adding image {image_path}: {type(e).__name__} - {e}. Skipping.", file=sys.stderr)
                p = doc.add_paragraph(); add_runs_for_formatted_text(p, f"[Error adding image: {image_path}]")
    doc.save(output_docx_path)

# --- Internal Tests ---
def run_internal_parser_tests():
    print("Running Internal Parser Tests...")
    passed_count = 0
    failed_count = 0
    total_tests = 0
    # Define test cases directly - using raw strings for latex_input
    tests = [
        {"name": "test_parse_title", "latex_input": r"\title{My Test Title}", "expected_json": [{'type': 'title', 'content': 'My Test Title'}]},
        {"name": "test_parse_section", "latex_input": r"\section{Section One}", "expected_json": [{'type': 'heading1', 'content': 'Section One'}]},
        {"name": "test_parse_subsection", "latex_input": r"\subsection{Subsection A}", "expected_json": [{'type': 'heading2', 'content': 'Subsection A'}]},
        {"name": "test_parse_simple_paragraph_exact", "latex_input": "Just a plain paragraph.", "expected_json": [{'type': 'paragraph', 'content': [{'type': 'text', 'value': 'Just a plain paragraph.'}]}]},
        {"name": "test_parse_paragraph_with_bold", "latex_input": r"Some \textbf{bold} text.", "expected_json": [{'type': 'paragraph', 'content': [{'type': 'text', 'value': 'Some '}, {'type': 'text', 'value': '**bold**'}, {'type': 'text', 'value': ' text.'}]}]},
        {"name": "test_parse_paragraph_with_italic", "latex_input": r"An \textit{italic} word.", "expected_json": [{'type': 'paragraph', 'content': [{'type': 'text', 'value': 'An '}, {'type': 'text', 'value': '*italic*'}, {'type': 'text', 'value': ' word.'}]}]},
        {"name": "test_parse_itemize_list_simple", "latex_input": r"\begin{itemize}\item First item\item Second item\end{itemize}", "expected_json": [{'type': 'list_item_bullet', 'content': [{'type': 'text', 'value': 'First item'}]}, {'type': 'list_item_bullet', 'content': [{'type': 'text', 'value': 'Second item'}]}]},
        {"name": "test_parse_image_no_options", "latex_input": r"\includegraphics{images/my_pic.png}", "expected_json": [{'type': 'image', 'path': 'images/my_pic.png', 'options': {}}]},
        {"name": "test_parse_image_with_width_option", "latex_input": r"\includegraphics[width=10cm]{images/another_pic.jpeg}", "expected_json": [{'type': 'image', 'path': 'images/another_pic.jpeg', 'options': {'width': '10cm'}}]},
        {"name": "test_parse_image_with_textwidth_option", "latex_input": r"\includegraphics[width=0.5\textwidth]{images/test_pic.png}", "expected_json": [{'type': 'image', 'path': 'images/test_pic.png', 'options': {'width': '0.5\textwidth'}}]},
        {"name": "test_parse_citet_with_prenote", "latex_input": r"A citation \citet[see p.~15]{ref1} here.", "expected_json": [{'type': 'paragraph', 'content': [{'type': 'text', 'value': 'A citation '}, {'type': 'citation', 'command': 'citet', 'keys': ['ref1'], 'prenote': 'see p.~15', 'postnote': ''}, {'type': 'text', 'value': ' here.'}]}]},
        {"name": "test_parse_citep_multiple_keys", "latex_input": r"Another one \citep{ref2, ref3}.", "expected_json": [{'type': 'paragraph', 'content': [{'type': 'text', 'value': 'Another one '}, {'type': 'citation', 'command': 'citep', 'keys': ['ref2', 'ref3'], 'prenote': '', 'postnote': ''}, {'type': 'text', 'value': '.'}]}]}
    ]

    for test in tests:
        total_tests += 1
        test_name = test["name"]
        latex_input = test["latex_input"]
        expected_json = test["expected_json"]
        try:
            actual_json = parse_latex_to_json(latex_input)
            if actual_json == expected_json:
                print(f"[PASS] {test_name}")
                passed_count += 1
            else:
                print(f"[FAIL] {test_name}")
                print(f"  Input:    '{latex_input}'") # Show raw string input
                print(f"  Expected: {json.dumps(expected_json)}")
                print(f"  Actual:   {json.dumps(actual_json)}")
                failed_count += 1
        except Exception as e:
            print(f"[ERROR] {test_name} - Exception during test: {e}")
            failed_count += 1
            
    print("\nInternal Parser Test Summary:")
    print(f"Total tests: {total_tests}, Passed: {passed_count}, Failed: {failed_count}")
    if failed_count == 0: print("All internal parser tests passed!")
    else: print("Some internal parser tests failed.")

def main():
    parser = argparse.ArgumentParser(description='Convert LaTeX .tex file to .docx with Zotero citations.')
    parser.add_argument('input_file', nargs='?', default=None, help='Path to the input LaTeX (.tex) file (optional if running internal tests)') # type: ignore
    parser.add_argument('output_file', nargs='?', default=None, help='Path for the output DOCX (.docx) file (optional if running internal tests)') # type: ignore
    parser.add_argument('--run-internal-tests', action='store_true', help='Run embedded parser unit tests and exit.')
    args = parser.parse_args()

    if args.run_internal_tests:
        run_internal_parser_tests()
        sys.exit(0)

    if not args.input_file or not args.output_file:
        parser.print_help()
        print("\nError: input_file and output_file are required unless --run-internal-tests is specified.")
        sys.exit(1)

    try:
        print(f"Starting conversion of '{args.input_file}' to '{args.output_file}'...")
        with open(args.input_file, 'r', encoding='utf-8') as f: # type: ignore
            latex_text = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{args.input_file}' not found.", file=sys.stderr)
        sys.exit(1)
    except Exception as e: 
        print(f"Error reading input file '{args.input_file}': {e}", file=sys.stderr)
        sys.exit(1)

    try:
        parsed_json_data = parse_latex_to_json(latex_text)
        generate_docx(parsed_json_data, args.output_file) # type: ignore
        print(f"Conversion successful! Output written to '{args.output_file}'")
    except Exception as e:
        print(f"An error occurred during LaTeX parsing or DOCX generation: {e}", file=sys.stderr) # Changed this line
        sys.exit(1)

if __name__ == "__main__":
    main()
